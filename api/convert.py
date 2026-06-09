from http.server import BaseHTTPRequestHandler
import io, base64, os, re
from api._dxf_utils import parse_form_data, json_response, cors_preflight
import ezdxf


# ── SVG → DXF ──────────────────────────────────────────────────────────────

def svg_to_dxf(data):
    import xml.etree.ElementTree as ET
    doc = ezdxf.new('R2010')
    msp = doc.modelspace()
    root = ET.fromstring(data.decode('utf-8', errors='replace'))
    _svg_elem(root, msp)
    buf = io.BytesIO(); doc.saveas(buf); return buf.getvalue()

def _svg_elem(elem, msp):
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    try:
        if tag == 'line':
            msp.add_line((float(elem.get('x1',0)), -float(elem.get('y1',0))),
                         (float(elem.get('x2',0)), -float(elem.get('y2',0))))
        elif tag == 'rect':
            x,y,w,h = float(elem.get('x',0)),float(elem.get('y',0)),float(elem.get('width',0)),float(elem.get('height',0))
            msp.add_lwpolyline([(x,-y),(x+w,-y),(x+w,-(y+h)),(x,-(y+h)),(x,-y)])
        elif tag == 'circle':
            msp.add_circle((float(elem.get('cx',0)),-float(elem.get('cy',0))),float(elem.get('r',0)))
        elif tag == 'ellipse':
            cx,cy,rx,ry = float(elem.get('cx',0)),float(elem.get('cy',0)),float(elem.get('rx',0)),float(elem.get('ry',0))
            msp.add_circle((cx,-cy),rx) if abs(rx-ry)<0.01 else msp.add_ellipse((cx,-cy),(rx,0),ry/rx if rx else 0)
        elif tag in ('polyline','polygon'):
            nums = re.findall(r'[-+]?\d*\.?\d+', elem.get('points',''))
            pts = [(float(nums[i]),-float(nums[i+1])) for i in range(0,len(nums)-1,2)]
            if tag=='polygon' and pts: pts.append(pts[0])
            if len(pts)>=2: msp.add_lwpolyline(pts)
        elif tag == 'path':
            for seg in _parse_path(elem.get('d','')):
                if len(seg)>=2: msp.add_lwpolyline(seg)
    except Exception:
        pass
    for child in elem:
        _svg_elem(child, msp)

def _parse_path(d):
    segs, cur, x, y, sx, sy = [], [], 0, 0, 0, 0
    toks = re.findall(r'[MLHVCSQTAZmlhvcsqtaz]|[-+]?\d*\.?\d+(?:[eE][-+]?\d+)?', d)
    i, cmd = 0, 'M'
    while i < len(toks):
        t = toks[i]
        if t.isalpha():
            cmd = t; i += 1
            if cmd in 'Zz':
                if cur: cur.append((sx,-sy)); segs.append(cur); cur = []
            continue
        try:
            if cmd=='M':   x,y=float(toks[i]),float(toks[i+1]); sx,sy=x,y; cur and segs.append(cur); cur=[(x,-y)]; i+=2; cmd='L'
            elif cmd=='m': x+=float(toks[i]); y+=float(toks[i+1]); sx,sy=x,y; cur and segs.append(cur); cur=[(x,-y)]; i+=2; cmd='l'
            elif cmd=='L': x,y=float(toks[i]),float(toks[i+1]); cur.append((x,-y)); i+=2
            elif cmd=='l': x+=float(toks[i]); y+=float(toks[i+1]); cur.append((x,-y)); i+=2
            elif cmd=='H': x=float(toks[i]); cur.append((x,-y)); i+=1
            elif cmd=='h': x+=float(toks[i]); cur.append((x,-y)); i+=1
            elif cmd=='V': y=float(toks[i]); cur.append((x,-y)); i+=1
            elif cmd=='v': y+=float(toks[i]); cur.append((x,-y)); i+=1
            else: i+=1
        except (IndexError, ValueError): i+=1
    if cur and len(cur)>1: segs.append(cur)
    return segs


# ── IMAGE → DXF (raster edge trace) ────────────────────────────────────────

def image_to_dxf(data):
    from PIL import Image, ImageFilter
    img = Image.open(io.BytesIO(data)).convert('L')
    img = img.filter(ImageFilter.GaussianBlur(1.5))
    edges = img.filter(ImageFilter.FIND_EDGES)
    w, h = edges.size
    pix = edges.load()
    scale = 100.0 / max(w, h, 1)
    doc = ezdxf.new('R2010'); msp = doc.modelspace()
    for y in range(h):
        s = None
        for x in range(w):
            on = pix[x,y] > 40
            if on and s is None: s = x
            elif not on and s is not None:
                if x-s >= 2: msp.add_line((s*scale,(h-y)*scale),(x*scale,(h-y)*scale))
                s = None
        if s is not None and w-s>=2: msp.add_line((s*scale,(h-y)*scale),(w*scale,(h-y)*scale))
    for x in range(w):
        s = None
        for y in range(h):
            on = pix[x,y] > 40
            if on and s is None: s = y
            elif not on and s is not None:
                if y-s >= 2: msp.add_line((x*scale,(h-s)*scale),(x*scale,(h-y)*scale))
                s = None
    buf = io.BytesIO(); doc.saveas(buf); return buf.getvalue()


# ── TXT / CSV → DXF ────────────────────────────────────────────────────────

def txt_to_dxf(data):
    pts = []
    for line in data.decode('utf-8', errors='replace').splitlines():
        line = line.strip()
        if not line or line.startswith('#'): continue
        nums = re.findall(r'[-+]?\d*\.?\d+', line)
        if len(nums) >= 2: pts.append((float(nums[0]), float(nums[1])))
    doc = ezdxf.new('R2010'); msp = doc.modelspace()
    if len(pts) >= 2: msp.add_lwpolyline(pts)
    buf = io.BytesIO(); doc.saveas(buf); return buf.getvalue()


# ── XLSX → DXF ─────────────────────────────────────────────────────────────

def xlsx_to_dxf(data):
    import openpyxl
    ws = openpyxl.load_workbook(io.BytesIO(data)).active
    pts = []
    for row in ws.iter_rows(values_only=True):
        nums = [c for c in row if isinstance(c,(int,float))]
        if len(nums) >= 2: pts.append((float(nums[0]), float(nums[1])))
    doc = ezdxf.new('R2010'); msp = doc.modelspace()
    if len(pts) >= 2: msp.add_lwpolyline(pts)
    buf = io.BytesIO(); doc.saveas(buf); return buf.getvalue()


# ── DOCX → DXF ─────────────────────────────────────────────────────────────

def docx_to_dxf(data):
    from docx import Document
    d = Document(io.BytesIO(data))
    pts = []
    for table in d.tables:
        for row in table.rows:
            nums = []
            for cell in row.cells:
                nums += [float(n) for n in re.findall(r'[-+]?\d*\.?\d+', cell.text)]
            if len(nums) >= 2: pts.append((nums[0], nums[1]))
    if not pts:
        for para in d.paragraphs:
            nums = re.findall(r'[-+]?\d*\.?\d+', para.text)
            for i in range(0, len(nums)-1, 2): pts.append((float(nums[i]), float(nums[i+1])))
    doc = ezdxf.new('R2010'); msp = doc.modelspace()
    if len(pts) >= 2: msp.add_lwpolyline(pts)
    buf = io.BytesIO(); doc.saveas(buf); return buf.getvalue()


# ── ROUTING TABLE ──────────────────────────────────────────────────────────

DIRECT = {
    '.dxf':  (lambda d: d,    ''),
    '.svg':  (svg_to_dxf,     '_converted'),
    '.png':  (image_to_dxf,   '_traced'),
    '.jpg':  (image_to_dxf,   '_traced'),
    '.jpeg': (image_to_dxf,   '_traced'),
    '.bmp':  (image_to_dxf,   '_traced'),
    '.txt':  (txt_to_dxf,     '_coords'),
    '.csv':  (txt_to_dxf,     '_coords'),
    '.xlsx': (xlsx_to_dxf,    '_coords'),
    '.xls':  (xlsx_to_dxf,    '_coords'),
    '.docx': (docx_to_dxf,    '_coords'),
}

GUIDED = {
    '.dwg':    ('AutoCAD DWG',        ['Open in AutoCAD or LibreCAD', 'File → Save As → AutoCAD 2000 DXF (*.dxf)', 'Upload the .dxf here']),
    '.step':   ('STEP / STP',         ['Open in Fusion 360, FreeCAD, or SolidWorks', 'Select the face or sketch to export', 'File → Export → DXF (2D)', 'Upload the .dxf here']),
    '.stp':    ('STEP / STP',         ['Open in Fusion 360, FreeCAD, or SolidWorks', 'Select the face or sketch to export', 'File → Export → DXF (2D)', 'Upload the .dxf here']),
    '.f3d':    ('Fusion 360 (.f3d)',   ['Open in Autodesk Fusion 360', 'Right-click the sketch or flat pattern in the browser tree', 'Select Save As DXF', 'Upload the .dxf here']),
    '.fd3':    ('Fusion 360 (.fd3)',   ['Open in Autodesk Fusion 360', 'Right-click the sketch or flat pattern in the browser tree', 'Select Save As DXF', 'Upload the .dxf here']),
    '.sldprt': ('SolidWorks Part',    ['Open in SolidWorks', 'File → Save As → DXF/DWG', 'Select DXF output and choose the face/sketch', 'Upload the .dxf here']),
    '.sldasm': ('SolidWorks Assembly',['Open in SolidWorks', 'Create a drawing view of the flat pattern', 'File → Save As → DXF/DWG', 'Upload the .dxf here']),
    '.pdf':    ('PDF',                 ['Open in Inkscape (free download at inkscape.org)', 'File → Save As → Desktop Cutting Plotter (AutoCAD DXF R14)', 'Upload the .dxf here', 'Note: only works for PDFs with vector geometry, not scanned drawings']),
}


class handler(BaseHTTPRequestHandler):
    def do_GET(self):
        json_response(self, {
            'direct_support': list(DIRECT.keys()),
            'guided_export': {k: v[0] for k, v in GUIDED.items()},
        })

    def do_POST(self):
        try:
            fields, files = parse_form_data(self)
            fi = files.get('file')
            if not fi:
                json_response(self, {'error': 'No file uploaded'}, 400); return

            name = fi['filename']
            ext = os.path.splitext(name)[1].lower()
            data = fi['data']

            if ext in GUIDED:
                fmt, steps = GUIDED[ext]
                json_response(self, {'status': 'guided_export', 'format': fmt, 'instructions': steps}); return

            if ext in DIRECT:
                fn, suffix = DIRECT[ext]
                dxf = fn(data)
                base = os.path.splitext(name)[0]
                out = base + suffix + '.dxf'
                json_response(self, {
                    'status': 'converted',
                    'original_filename': name,
                    'output_filename': out,
                    'dxf_b64': base64.b64encode(dxf).decode(),
                    'size_bytes': len(dxf),
                }); return

            json_response(self, {
                'error': f'Unsupported format: {ext}',
                'direct_support': list(DIRECT.keys()),
                'guided_export': list(GUIDED.keys()),
            }, 400)

        except Exception as e:
            json_response(self, {'error': str(e)}, 500)

    def do_OPTIONS(self):
        cors_preflight(self)
